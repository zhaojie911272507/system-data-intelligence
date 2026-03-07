#!/usr/bin/env python3
"""
Universal document parser — auto-detects format and loads data.
Supports: .xlsx .xls .xlsm .et .docx .doc .wps .txt .md .rtz .csv .json
Usage: python doc_parser.py <filepath>
"""

import sys
import json
import re
import logging
from pathlib import Path

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


def load_excel(filepath: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    result = {ws.title: [list(r) for r in ws.iter_rows(values_only=True)] for ws in wb.worksheets}
    wb.close()
    return {'type': 'spreadsheet', 'sheets': result}


def load_excel_legacy(filepath: str) -> dict:
    import xlrd
    wb = xlrd.open_workbook(filepath)
    result = {ws.name: [ws.row_values(i) for i in range(ws.nrows)] for ws in wb.sheets()}
    return {'type': 'spreadsheet', 'sheets': result}


def load_word(filepath: str) -> dict:
    from docx import Document
    doc = Document(filepath)
    paragraphs = [{'style': p.style.name, 'text': p.text} for p in doc.paragraphs if p.text]
    tables = [[[cell.text for cell in row.cells] for row in t.rows] for t in doc.tables]
    return {'type': 'document', 'paragraphs': paragraphs, 'tables': tables}


def load_text(filepath: str) -> dict:
    try:
        import chardet
        with open(filepath, 'rb') as f:
            enc = chardet.detect(f.read(10000))['encoding'] or 'utf-8'
    except ImportError:
        enc = 'utf-8-sig'
    with open(filepath, encoding=enc, errors='replace') as f:
        text = f.read()
    return {'type': 'text', 'encoding': enc, 'content': text}


def load_markdown(filepath: str) -> dict:
    result = load_text(filepath)
    text = result['content']
    headings = re.findall(r'^(#{1,6})\s+(.+)$', text, re.MULTILINE)
    result['type'] = 'markdown'
    result['headings'] = [{'level': len(h[0]), 'text': h[1]} for h in headings]
    return result


def load_rtz(filepath: str) -> dict:
    import xml.etree.ElementTree as ET
    tree = ET.parse(filepath)
    root = tree.getroot()
    tag = root.tag
    ns = tag[1:tag.index('}')] if tag.startswith('{') else ''
    prefix = f'{{{ns}}}' if ns else ''

    def find(el, tag_name):
        return el.findtext(f'{prefix}{tag_name}') or ''

    tasks = []
    for task in root.iter(f'{prefix}Task'):
        if find(task, 'UID') == '0':
            continue
        tasks.append({
            'id': find(task, 'ID'),
            'name': find(task, 'Name'),
            'start': find(task, 'Start'),
            'finish': find(task, 'Finish'),
            'duration': find(task, 'Duration'),
            'percent_complete': find(task, 'PercentComplete'),
            'outline_level': find(task, 'OutlineLevel'),
        })
    return {'type': 'project', 'namespace': ns, 'tasks': tasks, 'task_count': len(tasks)}


def load_csv(filepath: str) -> dict:
    import pandas as pd
    df = pd.read_csv(filepath, sep=None, engine='python', encoding='utf-8-sig')
    return {'type': 'spreadsheet', 'sheets': {'Sheet1': df.values.tolist()}, 'columns': list(df.columns)}


def load_json(filepath: str) -> dict:
    with open(filepath, encoding='utf-8') as f:
        return {'type': 'json', 'data': json.load(f)}


LOADERS = {
    '.xlsx': load_excel, '.xlsm': load_excel, '.et': load_excel,
    '.xls':  load_excel_legacy,
    '.docx': load_word,  '.doc': load_word, '.wps': load_word,
    '.txt':  load_text,
    '.md':   load_markdown,
    '.rtz':  load_rtz,
    '.csv':  load_csv,
    '.json': load_json,
}


def detect_and_load(filepath: str) -> dict:
    """Auto-detect file format and load with the best strategy."""
    suffix = Path(filepath).suffix.lower()
    loader = LOADERS.get(suffix)
    if not loader:
        raise ValueError(f"Unsupported format: {suffix}. Supported: {list(LOADERS.keys())}")
    logger.info(f"Loading {suffix} with {loader.__name__}")
    return loader(filepath)


def main():
    if len(sys.argv) < 2:
        print("Usage: python doc_parser.py <filepath>")
        sys.exit(1)

    data = detect_and_load(sys.argv[1])
    output = json.dumps(data, ensure_ascii=False, default=str, indent=2)
    print(output)
    out_path = Path(sys.argv[1]).stem + "_parsed.json"
    Path(out_path).write_text(output, encoding='utf-8')
    logger.info(f"Saved to: {out_path}")


if __name__ == "__main__":
    main()
